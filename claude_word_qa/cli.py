import argparse
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from claude_word_qa.doc_parser import process_data_directory
from claude_word_qa.anthropic_client import ask_claude

def save_output(question: str, answer: str, technical_details: dict, output_dir: str = "output") -> str:
    """
    Save the question, answer, and technical details to a timestamped Word document in the output directory.
    Returns the filename of the saved file.
    """
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate timestamp for filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"claude_qa_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Claude Q&A Report', 0)
    
    # Add question section
    doc.add_heading('Question:', level=1)
    question_para = doc.add_paragraph(question)
    
    # Add separator
    doc.add_paragraph()
    
    # Add answer section
    doc.add_heading('Claude\'s Answer:', level=1)
    answer_para = doc.add_paragraph(answer)
    
    # Add separator and metadata
    doc.add_paragraph()
    doc.add_paragraph('_' * 60)
    metadata_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    # Apply italics formatting instead of style
    metadata_para.runs[0].italic = True
    
    # Add technical details section
    doc.add_paragraph()
    doc.add_heading('Technical Details:', level=1)
    
    # Model information
    doc.add_heading('LLM Model Used:', level=2)
    final_model = technical_details.get('final_model', 'Unknown')
    doc.add_paragraph(f"Primary Model: {final_model}")
    
    if technical_details.get('models_used'):
        models_text = ", ".join(technical_details['models_used'])
        doc.add_paragraph(f"All Models Used: {models_text}")
    
    # Batching information
    doc.add_heading('Document Processing Details:', level=2)
    doc.add_paragraph(f"Original Document Size: {technical_details.get('original_document_size', 0):,} characters")
    doc.add_paragraph(f"Chunks Created: {technical_details.get('chunks_created', 0)}")
    doc.add_paragraph(f"Chunks Successfully Processed: {technical_details.get('chunks_processed', 0)}")
    doc.add_paragraph(f"Chunks Failed: {technical_details.get('chunks_failed', 0)}")
    
    if technical_details.get('synthesis_performed'):
        doc.add_paragraph("Response Synthesis: Yes (multiple chunks combined)")
    else:
        doc.add_paragraph("Response Synthesis: No (single chunk processed)")
    
    # Detailed chunk information
    if technical_details.get('chunk_details'):
        doc.add_heading('Chunk Processing Details:', level=2)
        for chunk_detail in technical_details['chunk_details']:
            status = "✓ Success" if chunk_detail.get('success') else "❌ Failed"
            chunk_text = f"Chunk {chunk_detail.get('chunk_number', 'N/A')}: {chunk_detail.get('chunk_size', 0):,} chars - {chunk_detail.get('model_used', 'Unknown')} - {status}"
            doc.add_paragraph(chunk_text)
    
    # Save the document
    doc.save(filepath)
    
    return filename

def main():
    parser = argparse.ArgumentParser(
        description="Ask a question about the contents of documents in the data directory using Claude 4 Sonnet API."
    )
    parser.add_argument("question", type=str, help="The question to ask about the documents.")
    parser.add_argument(
        "--data-dir", type=str, default="data",
        help="Path to the data directory (default: data)"
    )
    parser.add_argument(
        "--output-dir", type=str, default="output",
        help="Path to the output directory for saving results (default: output)"
    )
    args = parser.parse_args()

    try:
        print("Loading and parsing documents from data directory...")
        document_text = process_data_directory(args.data_dir)
        print(f"Documents loaded successfully. Content length: {len(document_text)} characters")
        
        print(f"Sending question to Claude 4 Sonnet...")
        result = ask_claude(args.question, document_text)
        
        if result and result[0]:  # Check if we got a valid answer
            answer, technical_details = result
            
            print("\nClaude's answer:")
            print("-" * 60)
            print(answer)
            print("-" * 60)
            
            # Print technical summary
            print(f"\nTechnical Summary:")
            print(f"- LLM Model: {technical_details.get('final_model', 'Unknown')}")
            print(f"- Document Size: {technical_details.get('original_document_size', 0):,} characters")
            print(f"- Chunks: {technical_details.get('chunks_created', 0)} created, {technical_details.get('chunks_processed', 0)} processed")
            if technical_details.get('synthesis_performed'):
                print(f"- Response Synthesis: Yes")
            
            # Save the output
            filename = save_output(args.question, answer, technical_details, args.output_dir)
            print(f"\nOutput saved to: {args.output_dir}/{filename}")
            
        else:
            print("No answer received from Claude.")
            
    except FileNotFoundError as e:
        print(f"Error: {e}")
    except ValueError as e:
        print(f"Error: {e}")
    except UnicodeDecodeError as e:
        print(f"Unicode/Encoding Error: {e}")
        print("This might be due to special characters in the document or API response.")
    except Exception as e:
        print(f"Unexpected error: {e}")
        import traceback
        traceback.print_exc()
    
    # Program automatically quits after completion
    print("\nProgram completed.")

if __name__ == "__main__":
    main() 